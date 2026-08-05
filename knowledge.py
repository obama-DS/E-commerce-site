"""
knowledge.py — Knowledge Base for the Obama Store assistant.

A self-contained, file-backed knowledge store used by the chatbot. Items are
stored in SQLite, automatically parsed from many input formats (plain text,
Markdown, PDF, Word, CSV, JSON, website URLs), broken into searchable chunks
and indexed so the assistant can retrieve the latest uploaded knowledge on
every message — no model retraining required.

The API of this module (create / update / delete / search / stats) is used
directly by app.py. All public methods are safe to call from background
threads (a write lock guards SQLite writes).
"""

from __future__ import annotations

import csv
import html
import io
import json
import math
import os
import re
import sqlite3
import threading
import urllib.request
import uuid
from html.parser import HTMLParser
from pathlib import Path
from typing import Dict, List, Optional

# --------------------------------------------------------------------------
# Configuration
# --------------------------------------------------------------------------

BASE_DIR = Path(__file__).resolve().parent
DB_PATH = Path(os.environ.get("KB_DB_PATH", str(BASE_DIR / "knowledge.db")))

MAX_URL_BYTES = 2_000_000        # cap a downloaded page at ~2 MB
MAX_URL_TEXT = 100_000           # cap extracted page text at ~100k chars
MAX_ITEM_TEXT = 200_000          # cap any stored/imported text at ~200k chars
MAX_FILE_BYTES = 10_000_000      # cap an uploaded file at ~10 MB
MAX_JSON_CHARS = 60_000          # cap flattened JSON at ~60k chars
CSV_ROW_LIMIT = 600              # cap CSV rows that get indexed
CHUNK_SIZE = 900                 # target characters per indexed chunk
CHUNK_OVERLAP = 140              # overlap between consecutive chunks
TOKEN_MIN = 3                    # minimum token length that is indexed

_USER_AGENT = ("Mozilla/5.0 (compatible; ObamaStoreBot/1.0; "
               "+https://obamastore.example.com)")

# --------------------------------------------------------------------------
# Text helpers
# --------------------------------------------------------------------------


def _decode(data: bytes, content_type: str = "") -> str:
    """Decode raw bytes to text, trying the declared charset first."""
    charset = None
    match = re.search(r"charset=([\w-]+)", content_type or "", re.I)
    if match:
        charset = match.group(1)
    candidates = [charset, "utf-8", "utf-16", "windows-1252", "latin-1"]
    for enc in candidates:
        if not enc:
            continue
        try:
            return data.decode(enc)
        except (UnicodeDecodeError, LookupError):
            continue
    return data.decode("utf-8", errors="replace")


def _tokens(text: str) -> List[str]:
    """Lowercased alphanumeric words of meaningful length."""
    return re.findall(r"[a-z0-9]{3,}", (text or "").lower())


def _md_to_text(md: str) -> str:
    """Strip Markdown syntax for indexing (original Markdown is preserved)."""
    md = re.sub(r"```.*?```", " ", md, flags=re.S)
    md = re.sub(r"`([^`]*)`", r"\1", md)
    md = re.sub(r"!\[[^\]]*\]\([^)]*\)", " ", md)
    md = re.sub(r"\[([^\]]*)\]\([^)]*\)", r"\1", md)
    md = re.sub(r"(^|\n)\s*#{1,6}\s*", r"\1", md)
    md = re.sub(r"(^|\n)\s*[-*+]\s+", r"\1", md)
    md = re.sub(r"(^|\n)\s*\d+[.)]\s+", r"\1", md)
    md = re.sub(r"\*{1,3}|_{1,3}", "", md)
    md = re.sub(r"\n{3,}", "\n\n", md)
    return md.strip()


def _chunk_text(text: str, size: int = CHUNK_SIZE,
                overlap: int = CHUNK_OVERLAP) -> List[str]:
    """Split text into overlapping chunks along sentence boundaries."""
    text = (text or "").strip()
    if not text:
        return []
    sentences = re.split(r"(?<=[.!?])\s+|\n+", text)
    chunks: List[str] = []
    current = ""
    for sentence in sentences:
        sentence = sentence.strip()
        if not sentence:
            continue
        if len(current) + len(sentence) + 1 <= size or not current:
            current = (current + " " + sentence).strip()
        else:
            chunks.append(current)
            tail = " ".join(current.split()[-40:])
            current = (tail + " " + sentence).strip()
    if current:
        chunks.append(current)
    return [c for c in chunks if len(_tokens(c)) >= 2]


def _flatten_json(obj, prefix: str = "", depth: int = 0,
                  out: Optional[List[str]] = None) -> List[str]:
    if out is None:
        out = []
    if depth > 5:
        return out
    if isinstance(obj, dict):
        for key, value in obj.items():
            full = f"{prefix}.{key}" if prefix else str(key)
            if isinstance(value, (dict, list)):
                _flatten_json(value, full, depth + 1, out)
            else:
                out.append(f"{full}: {value}")
    elif isinstance(obj, list):
        for index, value in enumerate(obj[:60]):
            full = f"{prefix}[{index}]" if prefix else f"[{index}]"
            if isinstance(value, (dict, list)):
                _flatten_json(value, full, depth + 1, out)
            else:
                out.append(f"{full}: {value}")
    return out


# --------------------------------------------------------------------------
# Content extractors (file / URL -> plain text)
# --------------------------------------------------------------------------


def _extract_pdf(data: bytes) -> str:
    from pypdf import PdfReader  # installed separately
    reader = PdfReader(io.BytesIO(data))
    pages = [(page.extract_text() or "") for page in reader.pages]
    text = "\n\n".join(pages).strip()
    if not text:
        raise ValueError(
            "No extractable text found in the PDF. Scanned documents are not "
            "supported — please upload a text-based PDF.")
    return text


def _extract_word(data: bytes) -> str:
    from docx import Document  # installed separately
    doc = Document(io.BytesIO(data))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            parts.append(" | ".join(cells))
    text = "\n".join(parts).strip()
    if not text:
        raise ValueError("The Word document contains no readable text.")
    return text


def _extract_csv(data: bytes) -> str:
    text = _decode(data)
    reader = csv.reader(io.StringIO(text))
    rows = []
    for row in reader:
        if row and any(cell.strip() for cell in row):
            rows.append([cell.strip() for cell in row])
    if not rows:
        raise ValueError("The CSV file is empty.")
    header = rows[0]
    lines = [" | ".join(header)]
    for row in rows[1:CSV_ROW_LIMIT + 1]:
        lines.append(" | ".join(row))
    return "\n".join(lines)


def _extract_json(data: bytes) -> str:
    text = _decode(data)
    try:
        obj = json.loads(text)
    except json.JSONDecodeError as exc:
        raise ValueError(f"Invalid JSON: {exc.msg} at position {exc.pos}.") from exc
    lines = _flatten_json(obj)
    return "\n".join(lines)[:MAX_JSON_CHARS]


class _HtmlToText(HTMLParser):
    def __init__(self) -> None:
        super().__init__(convert_charrefs=True)
        self.parts: List[str] = []
        self.title = ""
        self._skip = 0
        self._in_title = False

    def handle_starttag(self, tag: str, attrs) -> None:
        if tag in ("script", "style", "noscript", "svg"):
            self._skip += 1
        if tag == "title":
            self._in_title = True
        if tag in ("p", "div", "br", "li", "tr", "section", "article",
                   "h1", "h2", "h3", "h4", "h5", "h6", "blockquote"):
            self.parts.append("\n")

    def handle_endtag(self, tag: str) -> None:
        if tag in ("script", "style", "noscript", "svg"):
            self._skip = max(0, self._skip - 1)
        if tag == "title":
            self._in_title = False

    def handle_data(self, data: str) -> None:
        if self._in_title and not self.title:
            self.title = data.strip()[:200]
        if self._skip == 0:
            self.parts.append(data)


def _html_to_text(raw: str):
    parser = _HtmlToText()
    try:
        parser.feed(raw[:1_500_000])
        parser.close()
    except Exception:
        pass
    lines = [ln.strip() for ln in "\n".join(parser.parts).split("\n")]
    text = "\n".join(ln for ln in lines if ln)
    return parser.title, text


def _fetch_url(url: str):
    """Download a URL and return (title, plain_text)."""
    request = urllib.request.Request(url, headers={"User-Agent": _USER_AGENT})
    with urllib.request.urlopen(request, timeout=20) as response:
        if getattr(response, "status", 200) != 200:
            raise ValueError(f"The URL responded with HTTP {response.status}.")
        content_type = response.headers.get("Content-Type", "")
        data = response.read(MAX_URL_BYTES)
    text = _decode(data, content_type)
    if "html" in content_type.lower() or "<html" in text[:2000].lower():
        title, text = _html_to_text(text)
    else:
        title = ""
        text = text.strip()
    if not text:
        raise ValueError("No readable text was found at this URL.")
    return title[:200], text[:MAX_URL_TEXT]


def _extract_file(filename: str, data: bytes) -> str:
    """Dispatch to the right extractor based on the file extension."""
    name = (filename or "").lower()
    if name.endswith(".pdf"):
        return _extract_pdf(data)
    if name.endswith(".docx"):
        return _extract_word(data)
    if name.endswith(".doc"):
        raise ValueError(
            "Legacy .doc files are not supported — please save the document "
            "as .docx and upload it again.")
    if name.endswith((".md", ".markdown")):
        return _decode(data)
    if name.endswith(".csv"):
        return _extract_csv(data)
    if name.endswith(".json"):
        return _extract_json(data)
    if name.endswith(".txt") or name.endswith(".text"):
        return _decode(data)
    raise ValueError(
        "Unsupported file type. Upload a PDF, Word (.docx), Markdown, CSV, "
        "JSON or text file.")


# --------------------------------------------------------------------------
# Knowledge Base
# --------------------------------------------------------------------------


class KnowledgeBase:
    def __init__(self, path: Path = DB_PATH) -> None:
        self.path = Path(path)
        self._write_lock = threading.Lock()
        self._init_db()

    # -- schema ---------------------------------------------------------

    def _init_db(self) -> None:
        self.path.parent.mkdir(parents=True, exist_ok=True)
        with self._conn() as conn:
            conn.executescript(
                """
                CREATE TABLE IF NOT EXISTS kb_items (
                    id            TEXT PRIMARY KEY,
                    title         TEXT NOT NULL,
                    content       TEXT NOT NULL,
                    content_type  TEXT NOT NULL DEFAULT 'text',
                    source        TEXT,
                    category      TEXT NOT NULL DEFAULT 'General',
                    tags          TEXT NOT NULL DEFAULT '[]',
                    status        TEXT NOT NULL DEFAULT 'processing',
                    error         TEXT,
                    char_count    INTEGER NOT NULL DEFAULT 0,
                    version       INTEGER NOT NULL DEFAULT 1,
                    created_at    INTEGER NOT NULL,
                    updated_at    INTEGER NOT NULL
                );
                CREATE TABLE IF NOT EXISTS kb_chunks (
                    id       INTEGER PRIMARY KEY AUTOINCREMENT,
                    item_id  TEXT NOT NULL,
                    seq      INTEGER NOT NULL,
                    content  TEXT NOT NULL
                );
                CREATE INDEX IF NOT EXISTS idx_chunks_item ON kb_chunks(item_id);
                CREATE TABLE IF NOT EXISTS kb_terms (
                    term      TEXT NOT NULL,
                    chunk_id  INTEGER NOT NULL,
                    count     INTEGER NOT NULL DEFAULT 1,
                    PRIMARY KEY (term, chunk_id)
                );
                CREATE INDEX IF NOT EXISTS idx_terms_term ON kb_terms(term);
                CREATE INDEX IF NOT EXISTS idx_items_status ON kb_items(status);
                CREATE INDEX IF NOT EXISTS idx_items_category ON kb_items(category);
                """
            )

    def _conn(self) -> sqlite3.Connection:
        conn = sqlite3.connect(self.path, timeout=30, check_same_thread=False)
        conn.row_factory = sqlite3.Row
        return conn

    # -- serialization helpers -----------------------------------------

    @staticmethod
    def _row_to_item(row: sqlite3.Row, include_content: bool = True) -> dict:
        if row is None:
            return {}
        item = dict(row)
        try:
            item["tags"] = json.loads(item.get("tags") or "[]")
        except (TypeError, ValueError):
            item["tags"] = []
        if not include_content:
            item["preview"] = (item.get("content") or "")[:200]
            item.pop("content", None)
        return item

    @staticmethod
    def _now() -> int:
        import time
        return int(time.time())

    # -- lifecycle ------------------------------------------------------

    def _set_status(self, item_id: str, status: str, error: str = "") -> None:
        with self._write_lock, self._conn() as conn:
            conn.execute(
                "UPDATE kb_items SET status = ?, error = ?, updated_at = ? "
                "WHERE id = ?",
                (status, error or None, self._now(), item_id),
            )

    def index_item(self, item_id: str) -> None:
        """(Re)parse, chunk and index a single item. Never raises."""
        self._set_status(item_id, "processing")
        try:
            item = self.get(item_id)
            if not item:
                return
            content = item["content"]
            plain = _md_to_text(content) if item["content_type"] == "markdown" \
                else content
            chunks = _chunk_text(plain)
            if not chunks:
                raise ValueError("No indexable text was found in this item.")
            terms: Dict[tuple, int] = {}
            with self._write_lock, self._conn() as conn:
                conn.execute("DELETE FROM kb_chunks WHERE item_id = ?",
                             (item_id,))
                conn.execute(
                    "DELETE FROM kb_terms WHERE chunk_id IN "
                    "(SELECT id FROM kb_chunks WHERE item_id = ?)",
                    (item_id,),
                )
                for seq, chunk in enumerate(chunks):
                    cursor = conn.execute(
                        "INSERT INTO kb_chunks (item_id, seq, content) "
                        "VALUES (?, ?, ?)",
                        (item_id, seq, chunk),
                    )
                    chunk_id = cursor.lastrowid
                    for term in _tokens(chunk):
                        terms[(term, chunk_id)] = \
                            terms.get((term, chunk_id), 0) + 1
                conn.executemany(
                    "INSERT INTO kb_terms (term, chunk_id, count) "
                    "VALUES (?, ?, ?) ON CONFLICT(term, chunk_id) "
                    "DO UPDATE SET count = excluded.count",
                    [(term, chunk_id, count)
                     for (term, chunk_id), count in terms.items()],
                )
                conn.execute(
                    "UPDATE kb_items SET char_count = ?, status = 'indexed', "
                    "error = NULL, updated_at = ? WHERE id = ?",
                    (len(plain), self._now(), item_id),
                )
        except Exception as exc:
            self._set_status(item_id, "error", str(exc)[:500])

    # -- CRUD -----------------------------------------------------------

    def create(self, title: str, content: str, category: str = "General",
               tags: Optional[List[str]] = None,
               content_type: str = "text", source: str = "") -> dict:
        title = (title or "Untitled").strip()
        content = (content or "").strip()[:MAX_ITEM_TEXT]
        if not content:
            raise ValueError("Content cannot be empty.")
        item_id = uuid.uuid4().hex[:12]
        now = self._now()
        payload = {
            "id": item_id,
            "title": title,
            "content": content,
            "content_type": content_type or "text",
            "source": (source or "").strip() or None,
            "category": (category or "General").strip() or "General",
            "tags": json.dumps([t.strip() for t in (tags or []) if t.strip()]),
            "status": "processing",
            "error": None,
            "char_count": 0,
            "version": 1,
            "created_at": now,
            "updated_at": now,
        }
        with self._write_lock, self._conn() as conn:
            conn.execute(
                "INSERT INTO kb_items (id, title, content, content_type, "
                "source, category, tags, status, error, char_count, version, "
                "created_at, updated_at) VALUES (:id, :title, :content, "
                ":content_type, :source, :category, :tags, :status, :error, "
                ":char_count, :version, :created_at, :updated_at)",
                payload,
            )
        self.index_item(item_id)
        return self.get(item_id)

    def update(self, item_id: str, title: Optional[str] = None,
               content: Optional[str] = None, category: Optional[str] = None,
               tags: Optional[List[str]] = None,
               content_type: Optional[str] = None,
               source: Optional[str] = None) -> dict:
        item = self.get(item_id)
        if not item:
            raise ValueError("Knowledge item not found.")
        with self._write_lock, self._conn() as conn:
            conn.execute(
                "UPDATE kb_items SET version = version + 1, updated_at = ? "
                "WHERE id = ?",
                (self._now(), item_id),
            )
        fields = {}
        if title is not None:
            fields["title"] = (title or "Untitled").strip()
        if content is not None:
            content = (content or "").strip()[:MAX_ITEM_TEXT]
            if not content:
                raise ValueError("Content cannot be empty.")
            fields["content"] = content
        if category is not None:
            fields["category"] = (category or "General").strip() or "General"
        if tags is not None:
            fields["tags"] = json.dumps(
                [t.strip() for t in tags if t.strip()])
        if content_type is not None:
            fields["content_type"] = content_type or "text"
        if source is not None:
            fields["source"] = (source or "").strip() or None
        if fields:
            sets = ", ".join(f"{key} = ?" for key in fields)
            values = list(fields.values()) + [item_id]
            with self._write_lock, self._conn() as conn:
                conn.execute(f"UPDATE kb_items SET {sets} WHERE id = ?",
                             values)
        self.index_item(item_id)
        return self.get(item_id)

    def delete(self, item_id: str) -> None:
        with self._write_lock, self._conn() as conn:
            conn.execute("DELETE FROM kb_chunks WHERE item_id = ?", (item_id,))
            conn.execute(
                "DELETE FROM kb_terms WHERE chunk_id IN "
                "(SELECT id FROM kb_chunks WHERE item_id = ?)",
                (item_id,),
            )
            conn.execute("DELETE FROM kb_items WHERE id = ?", (item_id,))

    @staticmethod
    def _content_type_for(filename: str) -> Optional[str]:
        name = (filename or "").lower()
        if name.endswith(".pdf"):
            return "pdf"
        if name.endswith(".docx"):
            return "word"
        if name.endswith((".md", ".markdown")):
            return "markdown"
        if name.endswith(".csv"):
            return "csv"
        if name.endswith(".json"):
            return "json"
        if name.endswith((".txt", ".text")):
            return "text"
        return None

    def _insert_pending(self, title: str, content_type: str, source: str,
                        category: str, tags: Optional[List[str]]) -> str:
        item_id = uuid.uuid4().hex[:12]
        now = self._now()
        with self._write_lock, self._conn() as conn:
            conn.execute(
                "INSERT INTO kb_items (id, title, content, content_type, "
                "source, category, tags, status, error, char_count, version, "
                "created_at, updated_at) VALUES (?, ?, '', ?, ?, ?, ?, "
                "'processing', NULL, 0, 1, ?, ?)",
                (item_id, (title or source or "Untitled").strip()[:200],
                 content_type, source, (category or "General").strip() or "General",
                 json.dumps([t.strip() for t in (tags or []) if t.strip()]),
                 now, now),
            )
        return item_id

    def _replace_content(self, item_id: str, content: str) -> None:
        with self._write_lock, self._conn() as conn:
            conn.execute(
                "UPDATE kb_items SET content = ?, updated_at = ? WHERE id = ?",
                (content[:MAX_ITEM_TEXT], self._now(), item_id),
            )

    def import_file(self, filename: str, data: bytes,
                    category: str = "General", tags: Optional[List[str]] = None,
                    title: str = "") -> dict:
        """Insert a file as a pending item, then parse + index in a thread."""
        ctype = self._content_type_for(filename)
        if not ctype:
            raise ValueError(
                "Unsupported file type. Upload a PDF, Word (.docx), Markdown, "
                "CSV, JSON or text file.")
        if len(data) > MAX_FILE_BYTES:
            raise ValueError("The file is too large (maximum 10 MB).")
        item_id = self._insert_pending(title, ctype, filename, category, tags)

        def _work() -> None:
            try:
                text = _extract_file(filename, data)
                self._replace_content(item_id, text)
                self.index_item(item_id)
            except Exception as exc:
                self._set_status(item_id, "error", str(exc)[:500])

        threading.Thread(target=_work, daemon=True).start()
        return {"id": item_id, "status": "processing", "content_type": ctype}

    def import_url(self, url: str, category: str = "General",
                   tags: Optional[List[str]] = None,
                   title: str = "") -> dict:
        """Insert a URL as a pending item, then download + index in a thread."""
        url = (url or "").strip()
        if not re.match(r"^https?://", url, re.I):
            raise ValueError("Enter a valid URL starting with http:// or https://.")
        item_id = self._insert_pending(title or url, "url", url, category, tags)

        def _work() -> None:
            try:
                page_title, text = _fetch_url(url)
                with self._write_lock, self._conn() as conn:
                    conn.execute(
                        "UPDATE kb_items SET content = ?, title = ?, "
                        "updated_at = ? WHERE id = ?",
                        (text, (title or page_title or url)[:200],
                         self._now(), item_id),
                    )
                self.index_item(item_id)
            except Exception as exc:
                self._set_status(item_id, "error", str(exc)[:500])

        threading.Thread(target=_work, daemon=True).start()
        return {"id": item_id, "status": "processing", "content_type": "url"}

    def get(self, item_id: str) -> dict:
        with self._conn() as conn:
            row = conn.execute(
                "SELECT * FROM kb_items WHERE id = ?", (item_id,)
            ).fetchone()
        return self._row_to_item(row) if row else {}

    def list(self, q: Optional[str] = None, category: Optional[str] = None,
             tag: Optional[str] = None, status: Optional[str] = None,
             page: int = 1, page_size: int = 20) -> dict:
        clauses = []
        params: List[str] = []
        if q:
            clauses.append(
                "(title LIKE ? OR content LIKE ?)")
            like = f"%{q}%"
            params += [like, like]
        if category:
            clauses.append("category = ?")
            params.append(category)
        if tag:
            clauses.append("tags LIKE ?")
            params.append(f"%\"{tag}\"%")
        if status:
            clauses.append("status = ?")
            params.append(status)
        where = f"WHERE {' AND '.join(clauses)}" if clauses else ""
        page = max(1, page)
        page_size = max(1, min(100, page_size or 20))
        offset = (page - 1) * page_size
        with self._conn() as conn:
            total = conn.execute(
                f"SELECT COUNT(*) FROM kb_items {where}", params
            ).fetchone()[0]
            rows = conn.execute(
                f"SELECT * FROM kb_items {where} ORDER BY updated_at DESC, "
                f"created_at DESC LIMIT ? OFFSET ?",
                params + [page_size, offset],
            ).fetchall()
        return {
            "total": total,
            "page": page,
            "page_size": page_size,
            "items": [self._row_to_item(r, include_content=False) for r in rows],
        }

    def stats(self) -> dict:
        with self._conn() as conn:
            rows = conn.execute(
                "SELECT status, COUNT(*) AS count FROM kb_items "
                "GROUP BY status"
            ).fetchall()
        counts = {r["status"]: r["count"] for r in rows}
        total = sum(counts.values())
        indexed = counts.get("indexed", 0)
        return {
            "total": total,
            "indexed": indexed,
            "processing": counts.get("processing", 0),
            "error": counts.get("error", 0),
            "categories": len(self.categories()),
            "coverage": round(indexed / total, 4) if total else 0.0,
        }

    def categories(self) -> List[dict]:
        with self._conn() as conn:
            rows = conn.execute(
                "SELECT category, COUNT(*) AS count FROM kb_items "
                "GROUP BY category ORDER BY count DESC, category"
            ).fetchall()
        return [{"name": r["category"], "count": r["count"]} for r in rows]

    def tags(self) -> List[dict]:
        seen: Dict[str, int] = {}
        with self._conn() as conn:
            rows = conn.execute("SELECT tags FROM kb_items").fetchall()
        for row in rows:
            try:
                for tag in json.loads(row["tags"] or "[]"):
                    seen[tag] = seen.get(tag, 0) + 1
            except (TypeError, ValueError):
                continue
        return [{"name": tag, "count": count}
                for tag, count in sorted(seen.items(),
                                         key=lambda kv: (-kv[1], kv[0]))]

    # -- search ---------------------------------------------------------

    def search(self, query: str, limit: int = 4) -> List[dict]:
        """Ranked chunk search. Returns chunks with idf-weighted scores.

        Content term matches earn idf-weighted points; query terms that also
        appear in the *item title* earn a strong bonus so that questions
        phrased around the topic of an entry (e.g. "return policy") rank it
        confidently above noise.
        """
        terms = _tokens(query)
        if not terms:
            return []
        with self._conn() as conn:
            total = conn.execute("SELECT COUNT(*) FROM kb_chunks").fetchone()[0]
            if total == 0:
                return []
            df_rows = conn.execute(
                "SELECT term, COUNT(*) AS df FROM kb_terms WHERE term IN "
                f"({','.join('?' * len(terms))}) GROUP BY term",
                terms,
            ).fetchall()
        df = {r["term"]: r["df"] for r in df_rows}
        idf = {t: math.log((total + 1) / (df.get(t, 0) + 1)) + 1.0 for t in terms}
        scores: Dict[int, float] = {}
        matched: Dict[int, set] = {}
        for term in terms:
            if term not in df:
                continue
            with self._conn() as conn:
                rows = conn.execute(
                    "SELECT chunk_id, count FROM kb_terms WHERE term = ?",
                    (term,),
                ).fetchall()
            for row in rows:
                chunk_id = row["chunk_id"]
                scores[chunk_id] = scores.get(chunk_id, 0.0) + \
                    idf[term] * math.sqrt(row["count"])
                matched.setdefault(chunk_id, set()).add(term)
        if not scores:
            return []
        # Title bonus: query terms found in the owning item's title.
        chunk_ids = list(scores.keys())
        placeholders = ",".join("?" * len(chunk_ids))
        with self._conn() as conn:
            chunk_rows = conn.execute(
                f"SELECT id, item_id FROM kb_chunks WHERE id IN ({placeholders})",
                chunk_ids,
            ).fetchall()
        item_ids = [r["item_id"] for r in chunk_rows]
        item_id_of = {r["id"]: r["item_id"] for r in chunk_rows}
        title_rows = conn.execute(
            f"SELECT id, title FROM kb_items WHERE id IN "
            f"({','.join('?' * len(item_ids))})",
            item_ids,
        ).fetchall()
        title_tokens = {r["id"]: set(_tokens(r["title"])) for r in title_rows}
        for chunk_id in chunk_ids:
            item_id = item_id_of.get(chunk_id)
            title_set = title_tokens.get(item_id, set()) if item_id else set()
            for term in matched.get(chunk_id, set()):
                if term in title_set:
                    scores[chunk_id] += idf[term] * 1.5
        top_ids = [cid for cid, _ in sorted(scores.items(),
                                            key=lambda kv: -kv[1])[:limit]]
        results = []
        for chunk_id in top_ids:
            with self._conn() as conn:
                row = conn.execute(
                    "SELECT c.item_id, c.id AS chunk_id, c.seq, c.content, "
                    "i.title, i.category, i.tags, i.content_type, i.source, "
                    "i.status, i.version, i.updated_at "
                    "FROM kb_chunks c JOIN kb_items i ON i.id = c.item_id "
                    "WHERE c.id = ?",
                    (chunk_id,),
                ).fetchone()
            if row and row["status"] == "indexed":
                item = self._row_to_item(row)
                item["score"] = round(scores[chunk_id], 3)
                results.append(item)
        return results

    def test(self, query: str, limit: int = 5) -> dict:
        results = self.search(query, limit=limit)
        return {
            "query": query,
            "results": results,
            "stats": self.stats(),
        }
